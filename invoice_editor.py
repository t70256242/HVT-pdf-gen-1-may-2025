# from docxtpl import DocxTemplate
# from jinja2 import Environment
#
#
# def sum_filter(iterable, attribute):
#     """Custom filter to sum values of a specific attribute in a list of dicts."""
#     return sum(float(item[attribute].replace(",", "")) for item in iterable)
#
#
# def invoice_edit(input_path, output_path, context):
#     # Load the template
#     doc = DocxTemplate(input_path)
#
#     # Create a custom Jinja2 environment and register the 'sum' filter
#     jinja_env = Environment()
#     jinja_env.filters['sum'] = sum_filter
#
#     # Render the template with the context and custom environment
#     doc.render(context, jinja_env)
#
#     # Save the filled document
#     doc.save(output_path)
#
#     print(f"{output_path} has been created!")
#
#
# # Example context with flexible payment descriptions and schedules
# context = {
#     # Client information
#     "date": "April 29, 2025",
#     "name": "Ojo Alaba",
#     "client_company_name": "Yoruba Ltd",
#     "client_phone": "+1 234 56000",
#     "client_address": "Lead Developer Street, Anthony, Riyah Turkey",
#     "client_email": "unfresh@email.com",
#     "project_name": "Tolu Scrapper",
#     "invoice_no": "#45678",
#
#     # Payment description table - flexible rows
#     "payment_description": [
#         {"s_no": "1", "description": "Project Setup Fee", "amount": "10,000"},
#         {"s_no": "2", "description": "Development Phase 1", "amount": "25,000"},
#         {"s_no": "3", "description": "API Integration", "amount": "15,000"},
#         {"s_no": "4", "description": "Project Setup Fee", "amount": "10,000"},
#         {"s_no": "5", "description": "Development Phase 1", "amount": "25,000"},
#         {"s_no": "5", "description": "API Integration", "amount": "15,000"}
#     ],
#
#     # Payment schedule table - flexible rows
#     "payment_schedule": [
#         {"s_no": "1", "schedule": "Upon signing", "amount": "10,000"},
#         {"s_no": "2", "schedule": "After 30 days", "amount": "25,000"},
#         {"s_no": "3", "schedule": "Final delivery", "amount": "15,000"},
#         {"s_no": "4", "schedule": "Upon signing", "amount": "10,000"},
#         {"s_no": "5", "schedule": "After 30 days", "amount": "25,000"},
#         {"s_no": "6", "schedule": "Final delivery", "amount": "15,000"},
#     ],
#
#     # # Calculate total automatically
#     # "total": "50,000",
# }
#
# invoice_edit("try_invoice_2_page_1.docx", "modified_invoice_2.docx", context)


from docxtpl import DocxTemplate
from jinja2 import Environment
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER


def sum_filter(iterable, attribute):
    """Custom filter to sum values of a specific attribute in a list of dicts."""
    return sum(float(item[attribute].replace(",", "")) for item in iterable)


def add_bold_line(paragraph, length=None, boldness=True, font_size=15):
    """Add a horizontal bold line using underscores, limited to page width."""
    # Calculate max length based on page width
    if length is None:
        content_width_in = 6.5  # Inches (8.5" page - 1" margins)
        char_width_in = 0.1  # Approximate width per character in inches
        max_chars = int(content_width_in / char_width_in)
        length = min(max_chars, 100)  # Use 100 as upper limit

    run = paragraph.add_run("_" * length)
    run.bold = boldness
    run.font.size = Pt(font_size)


# def add_payment_details_section(output_path, line_length=None, line_bold=True, line_font_size=15):
#     doc = Document(output_path)
#
#     # Add space before the top line
#     doc.add_paragraph()
#
#     # First bold line (auto-length)
#     top_line = doc.add_paragraph()
#     add_bold_line(top_line, length=line_length, boldness=line_bold, font_size=line_font_size)
#
#     # Space before the text
#     doc.add_paragraph()
#     doc.add_paragraph()
#
#     # "Payment Details:" heading
#     heading_para = doc.add_paragraph()
#     heading_run = heading_para.add_run("Payment Details:")
#     heading_run.bold = True
#     heading_run.font.size = Pt(15)
#     doc.add_paragraph()
#     doc.add_paragraph()
#
#     # Tab stop setup for alignment
#     tab_stops = heading_para.paragraph_format.tab_stops
#     tab_stops.add_tab_stop(Pt(200), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
#
#     # Payment detail lines with aligned columns
#     details = [
#         ("Name", "Hardik Vij"),
#         ("Account Holder", "HV TECHNOLOGIES"),
#         ("Account Number", "50200099469957"),
#         ("IFSC", "HDFC0000590"),
#         ("Branch", "GHAZIABAD - INDIRAPURAM"),
#         ("Account Type", "CURRENT"),
#     ]
#
#     for label, value in details:
#         para = doc.add_paragraph()
#         run = para.add_run(f"{label}:\t")
#         run.bold = True
#         run.font.size = Pt(11)
#         para.add_run(value).font.size = Pt(11)
#
#     # Space before the final line
#     doc.add_paragraph()
#
#     # Final bold line (auto-length)
#     bottom_line = doc.add_paragraph()
#     add_bold_line(bottom_line, length=line_length, boldness=line_bold, font_size=line_font_size)
#
#     doc.save(output_path)
#     print("Payment details section added.")

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_footer_to_docx(doc_path, y_offset_mm=15):
    """
    Adds a two-row left/right aligned footer with an adjustable vertical offset.

    Args:
        doc_path (str): Path to the .docx file.
        y_offset_mm (float): Distance from bottom of page to footer, in millimeters.
    """
    doc = Document(doc_path)
    section = doc.sections[0]
    footer = section.footer

    # Set vertical offset (from bottom of page)
    section.footer_distance = Pt(y_offset_mm * 2.835)  # mm to points

    # Clear existing footer paragraphs
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Calculate usable page width in points
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    usable_width = (page_width - left_margin - right_margin) / 12700  # EMU to pt

    def create_footer_line(left_text, right_text, right_color=RGBColor(0, 0, 0)):
        para = footer.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(Pt(usable_width), WD_ALIGN_PARAGRAPH.RIGHT)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Left side
        left_run = para.add_run(left_text)
        left_run.font.name = 'Calibri'
        left_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        left_run.font.size = Pt(11)
        left_run.font.color.rgb = RGBColor(0, 0, 0)  # Black

        para.add_run('\t')  # Tab to right side

        # Right side
        right_run = para.add_run(right_text)
        right_run.font.name = 'Calibri'
        right_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        right_run.font.size = Pt(11)
        right_run.font.color.rgb = right_color

    # Add two rows
    create_footer_line("+1 (469) 214-6349", "www.hvtechnologies.app", RGBColor(0, 102, 204))  # Blue link
    create_footer_line("+91-9773603818", "hardik@hvtechnologies.app")  # Black

    doc.save(doc_path)
    print(f"Footer added to {doc_path} with y-offset: {y_offset_mm}mm")





def set_footer(doc):
    for section in doc.sections:
        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Pt(500))
        table.autofit = True

        # Left column
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.text = "+91-9773603818\t\twww.hvtechnologies.app"

        # Right column
        right_cell = table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.text = "+91-8588099741\t\thardik@hvtechnologies.app"

def add_payment_details_section(output_path, line_length=None, line_bold=True, line_font_size=15):
    doc = Document(output_path)

    # Add footer to every page
    set_footer(doc)

    doc.add_paragraph()  # Add space before the top line

    # First bold line
    top_line = doc.add_paragraph()
    add_bold_line(top_line, length=line_length, boldness=line_bold, font_size=line_font_size)

    # Space before heading
    doc.add_paragraph()

    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run("Payment Details:")
    heading_run.bold = True
    heading_run.font.size = Pt(15)

    doc.add_paragraph()

    # Payment details
    details = [
        ("Name", "Hardik Vij"),
        ("Account Holder", "HV TECHNOLOGIES"),
        ("Account Number", "50200099469957"),
        ("IFSC", "HDFC0000590"),
        ("Branch", "GHAZIABAD - INDIRAPURAM"),
        ("Account Type", "CURRENT"),
    ]

    max_label_len = max(len(label) for label, _ in details)
    for label, value in details:
        para = doc.add_paragraph()
        label_text = f"{label}:{' ' * (max_label_len - len(label) + 2)}"
        run = para.add_run(label_text)
        run.bold = True
        run.font.size = Pt(11)
        val_run = para.add_run(value)
        val_run.font.size = Pt(11)

    doc.add_paragraph()  # Space before the final line

    # Final bold line
    bottom_line = doc.add_paragraph()
    add_bold_line(bottom_line, length=line_length, boldness=line_bold, font_size=line_font_size)

    # Space before T&Cs
    doc.add_paragraph()
    doc.add_paragraph()

    terms_heading = doc.add_paragraph()
    terms_heading_run = terms_heading.add_run("Terms and Conditions:")
    terms_heading_run.bold = True
    terms_heading_run.font.size = Pt(12)

    terms = [
        "Payment needs to be released as per the schedule mentioned in the proposal.",
        "Any out-of-scope work is subject to additional charges."
    ]

    for term in terms:
        para = doc.add_paragraph()
        run = para.add_run("•\t")
        run.font.size = Pt(11)
        term_run = para.add_run(term)
        term_run.font.size = Pt(11)

    doc.save(output_path)
    print("Payment details and terms section added with footer.")


# def invoice_edit(input_path, output_path, context):
#     doc = DocxTemplate(input_path)
#
#     # Custom Jinja2 env with sum filter
#     jinja_env = Environment()
#     jinja_env.filters['sum'] = sum_filter
#
#     # Render and save
#     doc.render(context, jinja_env)
#     doc.save(output_path)
#
#     print(f"{output_path} has been created!")


from docxtpl import DocxTemplate
from jinja2 import Environment
from num2words import num2words  # ✅ Import number-to-words converter

def sum_filter(values):
    return sum(values)

def invoice_edit(input_path, output_path, context):
    import re
    doc = DocxTemplate(input_path)

    # Sum payment_description prices
    # total_price = 0
    # for item in context["payment_description"]:
    #     # Extract numbers and decimals only, e.g. "USD 5,000.25" → "5000.25"
    #     cleaned = re.sub(r"[^\d.]", "", str(item["price"]))
    #
    #     # Handle empty or invalid price fallback
    #     if cleaned:
    #         try:
    #             numeric_price = float(cleaned)
    #         except ValueError:
    #             numeric_price = 0
    #     else:
    #         numeric_price = 0
    #
    #     # Update the item to have formatted price (e.g. "5,000.25")
    #     item["price"] = f"{numeric_price:,.2f}"
    #     total_price += numeric_price
    # # Add total sum and its word version
    # context["sum"] = f"{total_price:,}"
    # context["sum_to_word"] = num2words(total_price, to="cardinal", lang="en").title()  # e.g. "Fifty Thousand"

    # Render
    jinja_env = Environment()
    # jinja_env.filters['sum'] = sum_filter
    doc.render(context, jinja_env)
    doc.save(output_path)

    print(f"{output_path} has been created!")




# # Example context
# context = {
#     "date": "April 29, 2025",
#     "client_name": "Ojo Alaba",
#     "client_company_name": "Yoruba Ltd",
#     "client_phone": "+1 234 56000",
#     "company_number": "+1 234 56000",
#     "company_gst": "9000",
#     "client_address": "Lead Developer Street, Anthony, Riyah Turkey",
#     "client_email": "unfresh@email.com",
#     "project_name": "Tolu Scrapper",
#     "invoice_no": "#45678",
#     "payment_currency": "USD $",
#     "payment_description": [
#         {"s_no": "1", "description": "Project Setup Fee", "hns_code": 2345666, "price": "10,000"},
#         {"s_no": "2", "description": "Development Phase 1", "hns_code": 2345666, "price": "25,000"},
#         {"s_no": "3", "description": "API Integration", "hns_code": 2345666, "price": "15,000"},
#     ],
#     "payment_schedule": [
#         {"s_no": "1", "schedule": "Upon signing", "price": "10,000"},
#         {"s_no": "2", "schedule": "After 30 days", "price": "25,000"},
#         {"s_no": "3", "schedule": "Final delivery", "price": "15,000"},
#     ],
# }
#
# invoice_edit("gen_4.docx", "gen_invoice_1.docx", context)

