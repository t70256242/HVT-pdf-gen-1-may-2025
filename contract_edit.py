from docx import Document

def process_paragraph(paragraph, replacements):
    """Replace placeholders across all runs in a paragraph while preserving style outside placeholders."""
    # Combine all runs' text
    full_text = ''.join(run.text for run in paragraph.runs)

    # Replace placeholders
    for placeholder, value in replacements.items():
        full_text = full_text.replace(placeholder, value)

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''

    # Assign new text to the first run
    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_docx_placeholders(input_path, output_path, replacements):
    """
    Replace placeholders in a Word document (paragraphs and tables).
    Handles placeholders split across multiple runs.
    """
    doc = Document(input_path)

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, replacements)

    # Process all table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, replacements)

    doc.save(output_path)
    print(f"Successfully generated contract at: {output_path}")

# Example replacements
# replacements = {
#     "_Date_": "April 29, 2025",
#     "_Client Company Name_": "Acme Corporation",
#     "_Client Company Address_": "123 Business Rd, New York",
#     "_Contract End_": "April 29, 2026",
# }
#
# # Call function
# replace_docx_placeholders(
#     input_path="Contract Template_Format_.docx",
#     output_path="Generated_Contract_105.docx",
#     replacements=replacements
# )




