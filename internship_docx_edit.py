from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re


def replace_docx_placeholders(input_path, output_path, replacements):

    doc = Document(input_path)

    patterns = {
        "_Internship_Duration_": re.compile(r'(_Internship_Duration_)'),
        "_First_Pay_Cheque_Date": re.compile(r'(_First_Pay_Cheque_Date)'),
        "standard": re.compile(r'(_[^_]+_)')
    }

    def process_run(run):
        text = run.text

        for ph in ["_Internship_Duration_", "_First_Pay_Cheque_Date"]:
            if ph in replacements and ph in text:
                text = text.replace(ph, replacements[ph])

        matches = patterns["standard"].findall(text)
        for match in matches:
            if match in replacements:
                text = text.replace(match, replacements[match])

        run.text = text

    for paragraph in doc.paragraphs:

        style = paragraph.style

        for run in paragraph.runs:
            process_run(run)
        paragraph.style = style

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        process_run(run)

    doc.save(output_path)
    print(f"Document with preserved formatting saved to {output_path}")


# replacements = {
#     "_Date_": "April 29, 2025",
#     "_Name_": "John Doe",
#     "_Position_": "Software Engineering",
#     "_Stipend_": "15,000",
#     "_Hrs_": "20",
#     "_Internship_Duration_": "Three",
#     "_First_Pay_Cheque_Date": "May 1, 2025"
# }

# replace_docx_placeholders(
#     input_path="Internship Offer Letter Template.docx",
#     output_path="Modified_Offer_Letter_Formatted77.docx",
#     replacements=replacements
# )
